import pypandoc
from docx import Document

def atualizar_documento(file_path, secao, texto):
    try:
        # Converta ODT para DOCX
        output_file = file_path.replace('.odt', '.docx')
        pypandoc.convert_file(file_path, 'docx', outputfile=output_file)

        # Manipule o DOCX
        doc = Document(output_file)

        if secao == 'historico':
            doc.add_heading('Histórico do Veículo', level=1)
            doc.add_paragraph(texto)

        # Salve o DOCX e converta de volta para ODT
        doc.save(output_file)
        pypandoc.convert_file(output_file, 'odt', outputfile=file_path)
    except Exception as e:
        raise Exception(f"Erro ao atualizar o documento: {e}")
